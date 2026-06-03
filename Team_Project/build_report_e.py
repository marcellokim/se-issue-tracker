from __future__ import annotations

import struct
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_ROOT = ROOT / "docs" / "screenshots" / "report-ui-2026-06-02"
OUTPUT_DOCX = Path(__file__).resolve().with_name("E_구현결과.docx")
FONT_NAME = "Noto Sans CJK KR"


@dataclass(frozen=True)
class ScreenSpec:
    number: int
    slug: str
    title: str
    report_point: str
    description: str
    required_elements: str

    @property
    def filename(self) -> str:
        return f"{self.number:02d}-{self.slug}.png"


def screen_specs() -> list[ScreenSpec]:
    return [
        ScreenSpec(
            1,
            "login",
            "로그인",
            "시스템 실행 진입점, 계정별 로그인",
            "active 계정만 로그인 가능하며 실패 메시지를 표시한다.",
            "계정 입력 폼과 로그인 동작 진입점",
        ),
        ScreenSpec(
            2,
            "admin-dashboard",
            "Admin 대시보드",
            "Admin이 프로젝트/계정 관리자인 점",
            "Admin 로그인 직후 전체 프로젝트와 사용자 요약을 확인한다.",
            "프로젝트/사용자 수 요약",
        ),
        ScreenSpec(
            3,
            "project-participants",
            "프로젝트 참여자 관리",
            "프로젝트 기반 권한과 참여자 관리",
            "Admin이 프로젝트 상세에서 참여자와 역할을 관리한다.",
            "참여자 목록, 이름, 역할",
        ),
        ScreenSpec(
            4,
            "project-list",
            "프로젝트 목록",
            "Non-admin 사용자의 프로젝트 접근 범위",
            "PL/Dev/Tester는 자신이 active member인 프로젝트만 조회한다.",
            "참여 프로젝트 목록",
        ),
        ScreenSpec(
            5,
            "project-issue-list",
            "프로젝트 이슈 목록",
            "프로젝트 멤버의 일반 이슈 조회",
            "프로젝트 일반 이슈를 조회하고 등록일시 기준 정보와 새 issue id 형식을 확인한다.",
            "reported/Reported 등록일시와 새 issue id 형식",
        ),
        ScreenSpec(
            6,
            "register-issue-form",
            "이슈 등록",
            "UC1 Register Issue",
            "프로젝트 참여자가 새 이슈의 제목, 설명, 우선순위를 입력한다.",
            "Title, Description, Priority 입력 폼",
        ),
        ScreenSpec(
            7,
            "issue-detail",
            "이슈 상세",
            "UC4 View Issue Detail, 댓글/이력/가능 액션",
            "이슈 기본 정보, 담당자, 댓글, 변경 이력, 의존성, 가능한 action을 확인한다.",
            "Reported/Updated 일시, 댓글 날짜, History 항목",
        ),
        ScreenSpec(
            8,
            "assignment-form",
            "이슈 배정",
            "UC5 Assignment, 추천 후보와 전체 후보",
            "PL이 상태별 assignee/verifier 후보를 확인하고 추천 결과를 참고해 배정한다.",
            "[Recommended] 후보와 추천 근거",
        ),
        ScreenSpec(
            9,
            "status-change-form",
            "상태 변경",
            "UC6 Change Issue State, 사유 댓글",
            "Dev/Tester/PL이 역할과 현재 상태에 맞는 상태 전이를 사유 댓글과 함께 수행한다.",
            "전이 방향과 사유 입력란",
        ),
        ScreenSpec(
            10,
            "dependency-detail",
            "의존성 관리",
            "UC7 Manage Dependency",
            "PL이 같은 프로젝트 안의 blocking/blocked 관계를 확인하고 resolved guard에 활용한다.",
            "Blocked by/Blocking 관계와 History 항목",
        ),
        ScreenSpec(
            11,
            "deleted-issues",
            "삭제 이슈 관리",
            "UC9 Deleted Issue Management",
            "PL이 soft-deleted issue를 조회하고 restore 또는 purge 흐름을 수행한다.",
            "DELETED 목록과 보관 한도",
        ),
        ScreenSpec(
            12,
            "statistics",
            "통계",
            "UC10 Statistics, 상태/우선순위 분포",
            "프로젝트 기준 상태, 우선순위, 기간별 이슈 통계를 확인한다.",
            "JavaFX 차트와 Swing 표",
        ),
        ScreenSpec(
            13,
            "account-management",
            "계정 관리",
            "UC12 Account Management",
            "Admin이 계정 생성, 역할 변경, 활성화와 비활성화 처리를 수행한다.",
            "계정 목록과 생성/역할변경/활성화 버튼",
        ),
        ScreenSpec(
            14,
            "project-management",
            "프로젝트 관리",
            "UC13 Project Management",
            "Admin이 프로젝트 목록을 확인하고 새 프로젝트 생성을 시작한다.",
            "프로젝트 목록과 생성 버튼",
        ),
    ]


def png_size(path: Path) -> tuple[int, int]:
    with path.open("rb") as source:
        header = source.read(24)
    if len(header) < 24 or not header.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError(f"{path} is not a PNG file")
    return struct.unpack(">II", header[16:24])


def screenshot_path(toolkit: str, spec: ScreenSpec) -> Path:
    return SCREENSHOT_ROOT / toolkit / spec.filename


def validate_screenshots(specs: list[ScreenSpec]) -> None:
    missing: list[str] = []
    bad_size: list[str] = []
    for spec in specs:
        for toolkit in ("javafx", "swing"):
            path = screenshot_path(toolkit, spec)
            if not path.exists():
                missing.append(str(path))
                continue
            if png_size(path) != (1280, 900):
                bad_size.append(f"{path}: {png_size(path)}")
    if missing or bad_size:
        lines = []
        if missing:
            lines.append("Missing screenshots:")
            lines.extend(f"- {item}" for item in missing)
        if bad_size:
            lines.append("Screenshots with unexpected dimensions:")
            lines.extend(f"- {item}" for item in bad_size)
        raise FileNotFoundError("\n".join(lines))


def add_run(paragraph, text: str, *, bold: bool = False, size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    styles["Normal"].font.size = Pt(10)
    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        styles[style_name].font.name = FONT_NAME
        styles[style_name].element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        styles[style_name].font.size = Pt(size)


def add_intro(document: Document) -> None:
    document.add_heading("E. 구현 결과", level=1)
    document.add_heading("1. 실행 구조", level=2)
    document.add_paragraph(
        "기본 실행은 JavaFX UI이며, Swing UI는 별도 Gradle task로 실행한다. "
        "두 UI는 같은 controller, service, domain, repository 계층을 재사용한다."
    )
    commands = document.add_paragraph()
    add_run(commands, "./gradlew run --console=plain\n", bold=True)
    add_run(commands, "./gradlew runSwing --console=plain\n", bold=True)
    add_run(commands, "./gradlew oracleLocalResetFixedSeed --console=plain", bold=True)

    document.add_heading("2. 주요 화면과 기능", level=2)
    document.add_paragraph(
        "아래 화면은 fixed seed 데이터 기준으로 JavaFX와 Swing을 같은 순서와 파일명으로 캡처한 결과이다. "
        "이슈 목록과 상세 화면에는 새 issue id 형식, Reported/Updated 일시, 댓글 날짜, History, dependency 정보가 표시된다."
    )


def add_screen_page(document: Document, spec: ScreenSpec) -> None:
    document.add_heading(f"2.{spec.number:02d}. {spec.title}", level=3)

    summary = document.add_paragraph()
    add_run(summary, "보고서 연결: ", bold=True)
    add_run(summary, spec.report_point)
    summary.add_run("\n")
    add_run(summary, "구현 설명: ", bold=True)
    add_run(summary, spec.description)
    summary.add_run("\n")
    add_run(summary, "확인 요소: ", bold=True)
    add_run(summary, spec.required_elements)

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for column, toolkit in enumerate(("JavaFX", "Swing")):
        cell = table.cell(0, column)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1
        add_run(paragraph, toolkit, bold=True)

    for column, toolkit in enumerate(("javafx", "swing")):
        cell = table.cell(1, column)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1
        paragraph.add_run().add_picture(str(screenshot_path(toolkit, spec)), width=Inches(4.75))


def add_recommendation_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("3. Assignment 추천 기능", level=2)
    document.add_paragraph(
        "Assignment 추천 기능은 PL이 이슈를 배정할 때 의사결정을 돕기 위한 보조 기능이다. "
        "추천은 자동 배정이 아니라 후보 ranking을 제공하는 방식으로 설계하였다."
    )
    document.add_paragraph(
        "추천 후보는 현재 이슈의 title, description과 과거 해결 이력의 유사도를 기반으로 계산한다. "
        "resolved/closed 이슈에서 fixer/resolver로 참여한 이력과 프로젝트 active 후보를 함께 고려한다. "
        "데이터가 부족한 경우에도 전체 active 후보 목록을 제공하여 PL이 직접 배정할 수 있도록 한다."
    )


def build_document(output_path: Path = OUTPUT_DOCX) -> Path:
    specs = screen_specs()
    validate_screenshots(specs)

    document = Document()
    configure_document(document)
    add_intro(document)

    for index, spec in enumerate(specs):
        document.add_page_break()
        add_screen_page(document, spec)

    add_recommendation_section(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    output = build_document()
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
